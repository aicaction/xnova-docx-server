from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re

app = Flask(__name__)

# ── Colores xNova ──────────────────────────────────────────────
XNOVA_GREEN  = RGBColor(0x1A, 0x5C, 0x38)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY    = RGBColor(0x26, 0x26, 0x26)
LIGHT_GRAY   = RGBColor(0x88, 0x88, 0x88)
BORDER_COLOR = "1A5C38"
LIGHT_BORDER = "CCCCCC"

# ── Helpers XML ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_padding(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',str(top)),('bottom',str(bottom)),('left',str(left)),('right',str(right))]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), val)
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def add_horizontal_rule(doc, color=BORDER_COLOR):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def sp(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)

def add_inline_text(paragraph, text, bold=False, italic=False, color=None):
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = italic
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.bold = bold
        else:
            run = paragraph.add_run(part)
            run.bold = bold
            run.italic = italic
        if color:
            run.font.color.rgb = color
    return paragraph

def is_table_line(line):
    return '|' in line and line.strip().startswith('|') or line.count('|') >= 2

def is_separator_line(line):
    s = line.strip()
    return bool(re.match(r'^[\|\s\-:]+$', s)) and '-' in s

def build_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    rows = [r + [''] * (col_count - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_width = int(Cm(16) / col_count)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_width
            if r_idx == 0:
                set_cell_bg(cell, "1A5C38")
                set_cell_border(cell, "1A5C38", "4")
                set_cell_padding(cell, 60, 60, 100, 100)
                p = cell.paragraphs[0]
                p.clear()
                add_inline_text(p, cell_text, bold=True, color=WHITE)
                sp(p, 3, 3)
            else:
                bg = "F2F7F4" if r_idx % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
                set_cell_border(cell, LIGHT_BORDER, "4")
                set_cell_padding(cell, 60, 60, 100, 100)
                p = cell.paragraphs[0]
                p.clear()
                add_inline_text(p, cell_text)
                sp(p, 3, 3)

def add_warning_block(doc, title, content):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, "FFF8E7")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'right', 'bottom']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), 'E67E22')
    tcBorders.append(left)
    tcPr.append(tcBorders)
    set_cell_padding(cell, 80, 80, 160, 120)
    p_title = cell.paragraphs[0]
    p_title.clear()
    run = p_title.add_run(f"⚠ {title}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    sp(p_title, 0, 4)
    if content:
        p_body = cell.add_paragraph()
        add_inline_text(p_body, content)
        sp(p_body, 0, 4)
    doc.add_paragraph()

def parse_and_build_docx(text):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY

    lines = text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    warning_buffer = None
    warnings_section = False

    def flush_table():
        nonlocal in_table, table_rows
        if table_rows:
            build_table(doc, table_rows)
            doc.add_paragraph()
        in_table = False
        table_rows = []

    def flush_warning():
        nonlocal warning_buffer
        if warning_buffer:
            content = ' '.join(warning_buffer['lines']).strip()
            add_warning_block(doc, warning_buffer['title'], content)
            warning_buffer = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Línea vacía
        if not stripped:
            flush_table()
            flush_warning()
            i += 1
            continue

        # Separadores ---
        if re.match(r'^[-*]{3,}$', stripped):
            flush_table()
            flush_warning()
            add_horizontal_rule(doc)
            i += 1
            continue

        # BUSINESS CASE (cabecera)
        if re.match(r'^#*\s*BUSINESS CASE', stripped):
            flush_table()
            flush_warning()
            p = doc.add_paragraph()
            sp(p, 0, 2)
            run = p.add_run('BUSINESS CASE')
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = XNOVA_GREEN
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Subtítulo "Empresa x xNova"
        if re.search(r'[×xX].+xNova', stripped, re.IGNORECASE) and len(stripped) < 80:
            flush_table()
            flush_warning()
            p = doc.add_paragraph()
            sp(p, 0, 2)
            run = p.add_run(re.sub(r'^#+\s*', '', stripped))
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = DARK_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Preparado para...
        if re.match(r'^#*\s*Preparado para', stripped, re.IGNORECASE):
            flush_table()
            flush_warning()
            content = re.sub(r'^#+\s*', '', stripped)
            p = doc.add_paragraph()
            sp(p, 0, 14)
            run = p.add_run(content)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = LIGHT_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headers de sección numerados "## 1. TITULO" o "1. TITULO"
        section_match = re.match(r'^#{1,3}\s*(\d+)\.\s+(.+)$', stripped) or \
                        re.match(r'^(\d+)\.\s+([A-ZÁÉÍÓÚÑ\s\-—–()\w]+)$', stripped)
        if section_match:
            flush_table()
            flush_warning()
            num = section_match.group(1)
            title_text = section_match.group(2).strip().rstrip('#').strip()
            warnings_section = any(k in title_text.upper() for k in ['RIESGO', 'RISK', 'ADVERTENCIA'])
            p = doc.add_paragraph()
            sp(p, 14, 4)
            run = p.add_run(f"{num}. {title_text}")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = XNOVA_GREEN
            add_horizontal_rule(doc, BORDER_COLOR)
            i += 1
            continue

        # Header # (sin número)
        if stripped.startswith('# '):
            flush_table()
            flush_warning()
            content = stripped[2:].strip()
            p = doc.add_paragraph()
            sp(p, 8, 4)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = XNOVA_GREEN
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Header ##
        if stripped.startswith('## '):
            flush_table()
            flush_warning()
            content = stripped[3:].strip()
            p = doc.add_paragraph()
            sp(p, 8, 3)
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = XNOVA_GREEN
            i += 1
            continue

        # Tabla markdown
        if is_table_line(stripped):
            if is_separator_line(stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Bloque ⚠ advertencia (sección de riesgos)
        warn_match = re.match(r'^\*{0,2}⚠\s*(.+?)\*{0,2}$', stripped) or \
                     re.match(r'^\*{2}⚠\s*(.+?)\*{0,2}', stripped)
        if warnings_section and warn_match:
            flush_warning()
            title = warn_match.group(1).rstrip('*').strip()
            warning_buffer = {'title': title, 'lines': []}
            i += 1
            continue

        # Contenido de bloque de advertencia
        if warning_buffer is not None:
            if stripped:
                warning_buffer['lines'].append(stripped)
            i += 1
            continue

        # Bullets
        if re.match(r'^[-•*]\s+', stripped):
            content = re.sub(r'^[-•*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            sp(p, 1, 1)
            add_inline_text(p, content)
            i += 1
            continue

        # Pasos de escaleta "**Paso 01**" o "- **Paso"
        paso_match = re.match(r'^[-*]?\s*\*{0,2}(Paso\s+\d+[^:*]*)\*{0,2}[:\-–]\s*(.*)$', stripped)
        if paso_match:
            flush_warning()
            step_title = paso_match.group(1).strip()
            step_body = paso_match.group(2).strip()
            p = doc.add_paragraph()
            sp(p, 6, 2)
            r1 = p.add_run(f"→ {step_title}")
            r1.bold = True
            r1.font.color.rgb = XNOVA_GREEN
            if step_body:
                r2 = p.add_run(f": {step_body}")
                r2.bold = False
            i += 1
            continue

        # Nota en cursiva (* ... *)
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            content = stripped[1:-1].strip()
            p = doc.add_paragraph()
            sp(p, 3, 3)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Notas 💡 📌
        if stripped.startswith(('💡', '📌', '⚠')):
            p = doc.add_paragraph()
            sp(p, 4, 4)
            run = p.add_run(stripped)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        sp(p, 2, 3)
        add_inline_text(p, stripped)
        i += 1

    flush_table()
    flush_warning()

    return doc

@app.route('/health', methods=['GET'])
def health():
    return {'status': 'ok'}, 200

@app.route('/generate', methods=['POST'])
def generate():
    # Aceptar tanto JSON como form-urlencoded
    if request.content_type and 'application/json' in request.content_type:
        data = request.get_json(force=True, silent=True) or {}
    else:
        data = request.form.to_dict()
        if not data:
            data = request.get_json(force=True, silent=True) or {}

    text = data.get('text', '')
    filename = data.get('filename', 'BusinessCase_xNova.docx')

    if not text:
        return {'error': 'No text provided'}, 400

    doc = parse_and_build_docx(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
